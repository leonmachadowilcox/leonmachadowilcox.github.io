# Requires:
#   pip install python-docx beautifulsoup4 lxml
#
# Usage:
#   python build_self_assessment_and_update_index.py
#
# What it does:
#   - Writes self_assessment.md
#   - Writes self_assessment.docx (basic markdown -> docx)
#   - Updates index.html: replaces the inner HTML of <section id="self-assessment"> with the HTML block below
#   - Creates a timestamped backup of index.html before modifying

from datetime import datetime
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

# -----------------------------
# CONFIG: paths (edit if needed)
# -----------------------------
INDEX_PATH = Path("index.html")        # your portfolio index
MD_OUT     = Path("self_assessment.md")
DOCX_OUT   = Path("self_assessment.docx")

# -----------------------------
# SOURCE CONTENT (edit text only)
# -----------------------------
MARKDOWN = """# Professional Self-Assessment

I began my computer science degree at SNHU in **[start term/year]** and will complete it upon finishing this course in **[graduation term/year]**. Before enrolling, I brought a decade of professional experience in Quality Assurance, but my exposure to software development was primarily from a testing perspective. Throughout this program, I expanded my technical expertise to include full-stack development, algorithms and data structures, database design, and secure coding practices—preparing me for an SDET-style role.

## Core Strengths Beyond the Artifacts

- **Collaborating in Agile environments:** practiced sprint planning, Git workflows (branching, PRs), and reviews without disrupting teammates.
- **Communicating with stakeholders:** tailored explanations for technical peers and non-technical audiences (from CS-250 SDLC).
- **Security awareness:** from CS-405 Secure Coding—defensive programming, common vulnerability identification, and remediation.

## Enhancement One — Software Design & Engineering

**Artifact:** Java Appointment/Contact/Task Manager • **Origin:** CS-320 Software Test Automation  
Refactored loosely connected classes into a more modular design, centralized validation, enforced ID immutability, and expanded unit tests. Result: clearer code, better maintainability, and stronger foundation for automation.

## Enhancement Two — Algorithms & Data Structures

**Artifact:** C++ Course Planner • **Origin:** CS-260 Data Structures & Algorithms  
Added guarded input in the menu loop, deterministic sort-before-print, and defensive CSV parsing. Reinforced understanding of complexity/memory and safe data handling.

## Enhancement Three — Databases

**Artifact:** Flask/Dash Animal Shelter Dashboard • **Origin:** CS-340 Client/Server Development  
Connected to a database for CRUD, introduced role-based access, separated concerns (blueprints + services), and improved UI/UX for deployment readiness.

## Professional Growth

I evolved from a QA-focused professional into a developer-tester hybrid. My QA background sharpened testing and debugging, while SNHU coursework equipped me to design, build, and optimize systems end-to-end. That blend positions me well for SDET roles bridging development and quality engineering.
"""

# This is what will be injected into <section id="self-assessment">…</section> in index.html
SELF_ASSESSMENT_HTML = """
<div class="panel pad">
  <h2 style="margin-top:0">Professional Self-Assessment</h2>

  <p class="muted">
    I began my computer science degree at SNHU in <strong>Fall of 2022</strong> and will complete it upon finishing this course in
    <strong>Fall of 2025</strong>. Before enrolling, I brought a decade of professional experience in Quality Assurance, but my exposure
    to software development was primarily from a testing perspective. Throughout this program, I expanded my technical expertise to include
    full-stack development, algorithms and data structures, database design, and secure coding practices—preparing me for an SDET-style role.
  </p>

  <h3>Core Strengths Beyond the Artifacts</h3>
  <ul>
    <li><strong>Collaborating in Agile environments:</strong> practiced sprint planning, Git workflows (branching, PRs), and reviews without disrupting teammates.</li>
    <li><strong>Communicating with stakeholders:</strong> tailored explanations for technical peers and non-technical audiences (from CS-250 SDLC).</li>
    <li><strong>Security awareness:</strong> from CS-405 Secure Coding—defensive programming, common vulnerability identification, and remediation.</li>
  </ul>

  <h3>Enhancement One — Software Design & Engineering</h3>
  <p class="muted">
    <strong>Artifact:</strong> Java Appointment/Contact/Task Manager • <strong>Origin:</strong> CS-320 Software Test Automation<br />
    Refactored loosely connected classes into a more modular design, centralized validation, enforced ID immutability, and expanded unit tests.
    Result: clearer code, better maintainability, and stronger foundation for automation.
  </p>

  <h3>Enhancement Two — Algorithms & Data Structures</h3>
  <p class="muted">
    <strong>Artifact:</strong> C++ Course Planner • <strong>Origin:</strong> CS-260 Data Structures & Algorithms<br />
    Added guarded input in the menu loop, deterministic sort-before-print, and defensive CSV parsing. Reinforced understanding of
    complexity/memory and safe data handling.
  </p>

  <h3>Enhancement Three — Databases</h3>
  <p class="muted">
    <strong>Artifact:</strong> Flask/Dash Animal Shelter Dashboard • <strong>Origin:</strong> CS-340 Client/Server Development<br />
    Connected to a database for CRUD, introduced role-based access, separated concerns (blueprints + services), and improved UI/UX for deployment readiness.
  </p>

  <h3>Professional Growth</h3>
  <p class="muted">
    I evolved from a QA-focused professional into a developer-tester hybrid. My QA background sharpened testing and debugging, while SNHU coursework
    equipped me to design, build, and optimize systems end-to-end. That blend positions me well for SDET roles bridging development and quality engineering.
  </p>
</div>
""".strip()


# -----------------------------
# Helpers
# -----------------------------
def write_markdown(path: Path, text: str):
    path.write_text(text, encoding="utf-8")


def md_to_docx(md_text: str, out_path: Path):
    """Tiny Markdown -> DOCX (headings, bullets, paragraphs)."""
    doc = Document()
    # base font
    styles = doc.styles['Normal']
    styles.font.name = 'Calibri'
    styles.font.size = Pt(11)

    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()

        if not line:
            doc.add_paragraph("")  # blank line
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            continue

        # Bullets
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
            continue

        # Very light inline cleanup for DOCX
        cleaned = line.replace("**", "")
        cleaned = cleaned.replace("  ", " ")
        doc.add_paragraph(cleaned)

    doc.save(str(out_path))


def update_index_html(index_path: Path, injected_html: str) -> bool:
    """Replace inner HTML of <section id="self-assessment">…</section>."""
    if not index_path.exists():
        print(f"[!] {index_path} not found.")
        return False

    soup = BeautifulSoup(index_path.read_text(encoding="utf-8"), "lxml")

    target = soup.select_one("section#self-assessment")
    if target is None:
        # If the section doesn't exist, create it and append to <main>
        main = soup.select_one("main") or soup.body
        new_section = soup.new_tag("section", id="self-assessment")
        new_section.append(BeautifulSoup(injected_html, "lxml"))
        if main:
            main.append(new_section)
        else:
            soup.body.append(new_section)
    else:
        # Wipe existing children and replace with our new block
        target.clear()
        target.append(BeautifulSoup(injected_html, "lxml"))

    # Backup
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup = index_path.with_suffix(f".backup-{ts}.html")
    backup.write_text(soup.prettify(), encoding="utf-8")

    # Write updated (non-prettified to minimize diffs)
    index_path.write_text(str(soup), encoding="utf-8")
    return True


# -----------------------------
# Main
# -----------------------------
if __name__ == "__main__":
    # 1) Write Markdown
    write_markdown(MD_OUT, MARKDOWN)
    print(f"[✓] Wrote {MD_OUT}")

    # 2) Write DOCX
    md_to_docx(MARKDOWN, DOCX_OUT)
    print(f"[✓] Wrote {DOCX_OUT}")

    # 3) Update index.html
    ok = update_index_html(INDEX_PATH, SELF_ASSESSMENT_HTML)
    if ok:
        print(f"[✓] Updated {INDEX_PATH} (backup created next to it).")
    else:
        print("[!] Could not update index.html")